import tkinter as tk
from datetime import datetime
import os
from dotenv import load_dotenv
import docx
import homePage


def create_directory(dir_path):
    if not os.path.exists(dir_path):
        os.mkdir(path=dir_path)


def save_entry(data_dir, today_tmp, title, content):
    entry_title = title.get(1.0, tk.END).strip()
    entry_content = content.get(1.0, tk.END)

    entry_doc = docx.Document()

    entry_doc.add_heading(entry_title, 0)

    content_p = entry_doc.add_paragraph(f"\n{entry_content}")
    content_run = content_p.add_run()
    content_run.font.name = 'Comic Sans MS'
    content_run.font.size = 12

    entry_doc.save(f"{data_dir}\\{today_tmp}-{entry_title}.docx")


def add_new_entry(entry_title="", entry_content=""):

    load_dotenv()

    # Create the main window
    root = tk.Tk()
    root.title("Profess")
    root.geometry("600x700")  # Set the window size

    # Create a frame widget
    frame = tk.Frame(root, bg="lightblue", bd=5, relief=tk.FLAT)
    frame.pack(fill=tk.BOTH, expand=True)  # Pack the frame with padding

    today = datetime.today()
    today_title_format = today.strftime('%d %B %Y, %I:%M %p')

    title_lbl = tk.Label(frame, text=f"Date: {today_title_format}",
                         bg="lightblue", fg="royalblue",
                         font=("Harlow Solid Italic", 20))
    title_lbl.pack(pady=(20, 40))

    entry_title_label = tk.Label(frame, text=f"Title",
                                 bg="lightblue", fg="royalblue",
                                 font=("Comic Sans MS", 14, "bold"))
    entry_title_label.pack(anchor="w", padx=20, pady=(30, 0))
    entry_title_text = tk.Text(frame, height=2)
    entry_title_text.insert(tk.END, entry_title)
    entry_title_text.pack(fill=tk.X, expand=False, padx=20, pady=(0, 30))

    entry_text_label = tk.Label(frame, text=f"Pour out your heart...",
                                bg="lightblue", fg="royalblue",
                                font=("Comic Sans MS", 14, "bold"))
    entry_text_label.pack(anchor="w", padx=20)
    entry_text = tk.Text(frame, height=15)
    entry_text.pack(fill=tk.X, expand=False, padx=20, pady=(0, 30))
    entry_text.insert(tk.END, entry_content)

    today_directory_format = today.strftime('%Y%m%d')
    today_file_format = today.strftime('%Y%m%d-%H%M')
    journal_data_dir = os.getenv("JOURNAL_DATA")
    journal_data_today_dir = f"{journal_data_dir}Entries\\{today_directory_format}"

    save_button = tk.Button(frame,
                            text="Save",
                            font=("Comic Sans MS", 12),
                            command=lambda: [create_directory(f"{journal_data_dir}Entries"),
                                             create_directory(journal_data_today_dir),
                                             save_entry(journal_data_today_dir,
                                                        today_file_format,
                                                        entry_title_text,
                                                        entry_text)])
    save_button.pack(pady=(0, 10), padx=(20,0), side=tk.LEFT)

    save_and_close_button = tk.Button(frame,
                                      text="Save & Close",
                                      font=("Comic Sans MS", 12),
                                      command=lambda : [create_directory(f"{journal_data_dir}Entries"),
                                                        create_directory(journal_data_today_dir),
                                                        save_entry(journal_data_today_dir,
                                                                   today_file_format,
                                                                   entry_title_text,
                                                                   entry_text),
                                                        root.destroy(),
                                                        homePage.launch_home_page()]
                                      )
    save_and_close_button.pack(pady=(0, 10), padx=(30, 0), side=tk.LEFT)

    exit_button = tk.Button(frame,
                            text="Exit",
                            font=("Comic Sans MS", 12),
                            command=lambda: [root.destroy(), homePage.launch_home_page()])
    exit_button.pack(pady=(0, 10), padx=(30,0), side=tk.LEFT)

    # Start the main event loop
    root.mainloop()


if __name__ == '__main__':
    add_new_entry("asdads", "asdasdasdasdasd")
